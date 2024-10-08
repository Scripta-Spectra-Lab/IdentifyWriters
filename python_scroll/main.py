import pandas as pd
import numpy as np
import chardet
import re
import statistics
import math
import seaborn as sns
import matplotlib.pyplot as plt
from scipy.stats import t, binom, norm, ks_2samp, ttest_ind, permutation_test
import random
from docx import Document as WordDocument
import pickle
import warnings
warnings.filterwarnings("ignore")
import csv
import os, shutil
from pathlib import Path
import seaborn as sns
import matplotlib.pyplot as plt
from matplotlib.lines import Line2D
import altair as alt
from matplotlib.font_manager import FontProperties






class Python_Scroll:
    def __init__(self, file_path):
        ext = file_path.lower().split('.')[-1]

        if ext == 'txt':
            self.data = self.extract_text_from_txt(file_path)
        elif ext == 'docx':
            docx_text = self.extract_text_from_docx(file_path)
            self.data = pd.DataFrame({'Text': docx_text.split('\n')})
        else:
            raise ValueError("Unsupported file format")

    def extract_text_from_txt(self, file_path):
        with open(file_path, 'rb') as file:
            raw_data = file.read()
            result = chardet.detect(raw_data)
            encoding = result['encoding']
        return pd.read_csv(file_path, delimiter='\t', encoding=encoding, header=None)

    def extract_text_from_docx(self, docx_file):
        doc = WordDocument(docx_file)
        docx_text = []
        for paragraph in doc.paragraphs:
            docx_text.append(paragraph.text)
        return '\n'.join(docx_text)

#####################
### Preprocessing ###
#####################

# Data cleaning
    def clean_data(self, save=0):
        # Define a function to remove characters with big dots and the tab character
        def remove_big_dots_and_tab(text):
            cleaned_text = ''
            for char in text:
                if char == '֯' or char == '\u200e':
                    cleaned_text = cleaned_text[:-1]
                    continue
                else:
                    cleaned_text += char
            return cleaned_text

        # Remove characters with big dots and the following tab character (like ת֯\t)
        self.data = self.data.applymap(remove_big_dots_and_tab)
        
        # Clean the data by removing content between '^' symbols
        self.data = self.data.applymap(lambda x: x.replace('?', ''))  # Remove '?'
        self.data = self.data.applymap(lambda x: x.replace('.', ''))  # Remove '.
        self.data = self.data.applymap(lambda x: re.sub(r'\{{.*?\}}', '', str(x)))  # Remove {{***}}
        self.data = self.data.applymap(lambda x: re.sub(r'\⟦.*?\⟧', '', str(x))) # remove ⟦***⟧
        self.data = self.data.applymap(lambda x: re.sub(r'\[.*?\]', '', str(x))) # Remove [***]

        # Remove characters with little dots above them (like נׄ), keeping the character itself
        self.data = self.data.applymap(lambda x: re.sub(r'\u05C4', '', str(x)))

        #self.data = self.data.applymap(lambda x: x.replace('\xa0', ''))  # Remove '\xa0'
        self.data = self.data.applymap(lambda x: re.sub(r'\[.*?\]', '', str(x)))
        
        self.original_data = self.data.copy() # create a copy for corrections computation
        self.data = self.data.applymap(lambda x: re.sub(r'\^.*?\^', '', str(x))) # Remove ^***^

        # Create new index columns based on the specified format
        index_cols = self.data.iloc[:, -1].str.extract(r'(\d+):(\d+)', expand=True)
        new_index_1 = index_cols[0]
        new_index_2 = index_cols[1]

        # Remove the "*:*" index from the lines
        self.data.iloc[:, -1] = self.data.iloc[:, -1].str.replace(r'\d+:\d+', '', regex=True).str.strip()

        # Reset the index to integer values
        self.data = self.data.reset_index(drop=True)

        # Set the new indexes to the existing data table
        self.data.index = pd.MultiIndex.from_arrays([new_index_1, new_index_2])

        # Reset the column names
        self.data.columns = range(self.data.shape[1])

        # Save the cleaned data to an Excel file using 'utf-8' encoding
        if save == 1:    
            self.save_to_word()
    
    def sample_with_replacement(self, start_row, end_row, number_of_bootstraps):
        # Get the high-level index, convert to integers, sort, and convert back to strings
        high_index = sorted(self.data.index.levels[0], key=lambda x: int(x))
        
        # Check if the provided range is valid
        if start_row < 0 or end_row >= len(high_index):
            raise ValueError(f"The provided range ({start_row}, {end_row}) is out of bounds for the high-level index with length {len(high_index)}.")

        # Extract the high-level indexes within the specified range
        high_index_range = high_index[start_row-1:end_row]

        # Filter the DataFrame to only include rows with high-level indexes in the specified range
        filtered_df = self.data[self.data.index.get_level_values(0).isin(high_index_range)]

        # Perform sampling with replacement and concatenate the sampled rows
        sampled_dfs = []
        sampled_high_indexes = np.random.choice(high_index_range, size=number_of_bootstraps+1, replace=True)
        for i, high_index in enumerate(sampled_high_indexes, start=1):
            sampled_df = filtered_df.loc[[high_index]].copy()  # Ensure the high-level index is kept
            sampled_df.index = pd.MultiIndex.from_product([[i], sampled_df.index.get_level_values(1)])
            sampled_dfs.append(sampled_df)

        # Concatenate the sampled DataFrames to form a new multi-index DataFrame
        new_data = pd.concat(sampled_dfs)

        # Update self.data with the new sampled DataFrame
        self.data = new_data

        # Save the sampled DataFrame to an Excel file
        self.data.to_excel('bootstrap_dataframe_output_new_indexes.xlsx', index=True) 

    def save_to_word(self):
        document = WordDocument()

        # Iterate through the data and add each cell as a paragraph
        for i in range(self.data.shape[0]):
            row = self.data.iloc[i].tolist()
            for cell in row:
                if pd.notna(cell):
                    document.add_paragraph(str(cell))
                else:
                    document.add_paragraph("")  # Add an empty paragraph for NaN values

        # Save the document to a Word file
        document.save("Cleaned_data.docx")

    def load_data(self, path):
        self.t_data = pd.read_csv(path, delimiter='\t', header=None, encoding='latin-1')
        level_labels = ['Level1'] * len(self.data) + ['Level2'] * len(self.data)
        self.t_data.index = pd.MultiIndex.from_tuples(zip(level_labels, self.data.index))
        self.data = self.t_data

    def __str__(self):
        return str(self.data)

# Creating a new table of appearance of each word in plane and defective spelling

    def word_counts(self, word_triplets_list):
        # Initialize a new table to store the word counts
        self.word_counts_table = pd.DataFrame()

        for i, word_triplet in enumerate(word_triplets_list):
            max_chars = 0
            words = word_triplet[0]
            max_chars = word_triplet[1] if len(word_triplet) > 1 else 0
            allowed_chars = word_triplet[2] if len(word_triplet) > 2 else ()
            word1, word2 = words

            word1_column = f'{word1}'
            word2_column = f'{word2}'

            word1_counts = pd.DataFrame(columns=['col_1'], index=range(len(self.data.index.levels[0])))
            word2_counts = pd.DataFrame(columns=['col_2'], index=range(len(self.data.index.levels[0])))

            # Access data at the specified level using .xs
            for high_index in self.data.index.levels[0]:
                values = self.data.xs(high_index, level=0).squeeze().tolist()

                defective_spelling_count = 0
                plane_spelling_count = 0

                for value in values:
                    # Split value into words
                    values_words = [word.strip() for word in value.split()]
                    for i, value_word in enumerate(values_words):
                        value_word = value_word.replace('\u200F', '')
                        len_value_word = len(value_word)
                        if value_word.endswith(word1):
                            if len_value_word == len(word1):
                                defective_spelling_count += 1
                            # the checking word is larger but in a possible length
                            elif (len_value_word - len(word1) <= max_chars) and (len_value_word - len(word1) > 0):
                                if allowed_chars != ():
                                    if value_word.startswith(allowed_chars):
                                        defective_spelling_count += 1
                                else: # allowed_chars not specified, then all chars are premitted
                                    defective_spelling_count += 1
                                    
                        elif value_word.endswith(word2):
                            if len_value_word == len(word2):
                                plane_spelling_count += 1
                            # the checking word is larger but in a possible length
                            elif (len_value_word - len(word2) <= max_chars) and (len_value_word - len(word2) > 0):
                                if allowed_chars != ():
                                    if value_word.startswith(allowed_chars):
                                        plane_spelling_count += 1
                                else: # allowed_chars not specified, then all chars are premitted
                                    plane_spelling_count += 1
                                    
                word1_counts.iat[int(high_index) - 1, 0] = defective_spelling_count
                word2_counts.iat[int(high_index) - 1, 0] = plane_spelling_count

            self.word_counts_table[word1_column] = word1_counts
            self.word_counts_table[word2_column] = word2_counts

# update self.word_counts_table manually
    def update_word_counts_table(self, count_var_change):
        for var_change in count_var_change:
            if len(var_change) == 3:
                column_name, row_number, new_value = var_change
                try:
                    self.word_counts_table.at[int(row_number) - 1, column_name] = new_value
                except Exception as e:
                    print(f"Error updating table: {e}")

# Concate new rows to table of words count (self.data)
    def join_dataframes(self, file_path):
        df = pd.read_excel(file_path)
        self.word_counts_table = pd.concat([self.word_counts_table, df], axis=1)

# Create a table for the statistical tests resalts
    def create_tests_table(self, var_test=False, row_test=False, count_corrections=False):
        # checking how many more columns to add (var and num of row tests)
        boolean_vars = var_test + row_test + count_corrections
        length = (len(self.word_counts_table.columns) // 2 + boolean_vars)
        self.table = pd.DataFrame(columns=range(length))


###################################
# Calculation of auxiliary tables #
################################### 

######################
## row's length var ##
######################

# Functions for computing num of words and letters in a row
    def replace_with_letter_count(self): 
        # Create a new DataFrame with the same shape as the original table
        new_data = pd.DataFrame(index=self.data.index, columns=self.data.columns) #maybe here the problem

        # Function to count letters in a string (excluding spaces)
        def count_letters(text):
            return len(re.sub(r'\s', '', str(text)))

        # Iterate over each cell in the original table
        for row_idx, row in self.data.iterrows():
            for col_idx, cell in row.items():
                # Count the letters in the cell (excluding spaces)
                letter_count = count_letters(cell)
                # Assign the letter count to the corresponding cell in the new table
                new_data.loc[row_idx, col_idx] = letter_count

        self.letters_rows_count = new_data

    def compute_row_var(self):
        # Replace letters with their counts
        self.replace_with_letter_count()
        
        # Define the custom variance function
        def custom_variance(group):
            mean = group.mean()
            numerator = ((group - mean) ** 2) / mean
            variance = numerator.sum() / (len(group) - 1)
            return variance
        
        # Apply the custom variance function to each group
        group_stats = self.letters_rows_count.groupby(level=0).apply(custom_variance)
        group_stats.index = group_stats.index.astype(int)
        
        # Convert the resulting Series to a DataFrame if needed
        if isinstance(group_stats, pd.Series):
            group_stats = group_stats.to_frame(name='custom_var')
        
        # Sort the DataFrame based on the first index values
        self.var_table = group_stats.sort_index()

#####################
## rows in columns ##
#####################
    def compute_column_num_of_rows(self):
        # Calculate the mean and standard deviation for each group
        self.replace_with_letter_count()
        group_stats = self.letters_rows_count.groupby(level=0).agg(['count'])
        group_stats.index = group_stats.index.astype(int)
        
        # Sort the DataFrame based on the first index values
        self.num_of_rows_column = group_stats.sort_index()
        self.num_of_rows_column.at[self.num_of_rows_column.index[-1],0] = self.num_of_rows_column[self.num_of_rows_column.columns[0]].mean()
        # self.num_of_rows_column.at[27,0] = ''
        # self.num_of_rows_column.at[54,0] = ''
        
#####################
# count corrections #
#####################
    def count_corrections(self, char_corrections):
        
        # Apply the clean_text function to each element in the DataFrame
        self.data_t = self.data.copy()
        self.data = self.original_data.copy()
        
        # Create new index columns based on the specified format
        index_cols = self.data.iloc[:, -1].str.extract(r'(\d+):(\d+)', expand=True)
        new_index_1 = index_cols[0]
        new_index_2 = index_cols[1]

        # Remove the "*:*" index from the lines
        self.data.iloc[:, -1] = self.data.iloc[:, -1].str.replace(r'\d+:\d+', '', regex=True).str.strip()

        # Reset the index to integer values
        self.data = self.data.reset_index(drop=True)

        # Set the new indexes to the existing data table
        self.data.index = pd.MultiIndex.from_arrays([new_index_1, new_index_2])

        # Reset the column names
        self.data.columns = range(self.data.shape[1])
        
        self.corrections = self.count_num_of_corrections(char_corrections)
        self.data = self.data_t.copy()
        
    def count_num_of_corrections(self, char_corrections):                                     
        data = pd.DataFrame(columns=self.data.columns, index=range(1,len(self.data.index.levels[0])+1))

        for high_index in self.data.index.levels[0]:
            corrections_count = 0
            values = self.data.xs(high_index, level=0).squeeze().tolist()
            for item in values:
                words = item.split()
                for word in words:
                    count_carets = word.count(char_corrections)
                    if count_carets % 2 == 0:  # Check if count of '^' is even
                        corrections_count += count_carets // 2  # Add half the number of pairs to the total count
            data.iat[int(high_index) - 1, 0] = corrections_count
        return data


####################################
# Base Statistical Tests Functions #
####################################  
    # Auxiliary functions for the statistical tests

#### comparing same columns dist ####
    def read_col(self, col, start, stop):
        return self.rellavent_data[col].iloc[start-1:stop]
    def read_col_at_indices(self, col, indices):
        return self.rellavent_data[col].iloc[indices]
        
        
    def WaldTest_WithData(self, col_1_data_1, col_2_data_1, col_1_data_2, col_2_data_2):
        
        n_obs_1 = col_1_data_1.sum() + col_2_data_1.sum()
        n_obs_2 = col_1_data_2.sum() + col_2_data_2.sum()
        
        if n_obs_1 == 0 or n_obs_2 == 0:
            return ''            
        
        p_1 = col_1_data_1.sum()/n_obs_1
        p_2 = col_1_data_2.sum()/n_obs_2
        p_h = (n_obs_1*p_1 + n_obs_2*p_2) / (n_obs_1+n_obs_2)

        sd_h = math.sqrt(p_h*(1-p_h)*(1/n_obs_1 + 1/n_obs_2))
        if sd_h == 0:
            return ''
        
        degrees_of_freedom = n_obs_1 + n_obs_2 -2
        # if degrees_of_freedom + 2 >= 30:
        #   return self.update_p_value(norm.cdf((p_2 - p_1) / sd_h))
        # else:
        return self.update_p_value(t.cdf((p_2 - p_1) / sd_h, degrees_of_freedom))
        # return self.update_p_value(norm.cdf((p_2 - p_1) / sd_h))

    # Preparing the columns Wald to test
    def WaldTest(self, col_1, col_2, start_1, stop_1, start_2, stop_2):
        col_1_data_1 = self.read_col(col_1, start_1, stop_1)
        col_2_data_1 = self.read_col(col_2, start_1, stop_1)
        col_1_data_2 = self.read_col(col_1, start_2, stop_2)
        col_2_data_2 = self.read_col(col_2, start_2, stop_2)
        return self.WaldTest_WithData(col_1_data_1, col_2_data_1, col_1_data_2, col_2_data_2)


    def PermuteWald(self, col_1, col_2, start_1, stop_1, start_2, stop_2, t_value, n_iter = 1000):
        tests_results = []
        indexes_to_shuffle = np.concatenate((np.arange(start_1-1, stop_1), np.arange(start_2-1, stop_2)) )
        n_obs_1 = stop_1 - start_1
        #n_obs_2 = stop_2 - start_2
        for j in range(n_iter):
            if j%100 == 0:
                #print(f"Wald permutation test num: {j}")
                continue
            permuted_indices = np.random.permutation(indexes_to_shuffle)
            col_1_data = self.read_col_at_indices(col_1, permuted_indices)
            col_2_data = self.read_col_at_indices(col_2, permuted_indices)
            col_1_data_1 = col_1_data[:n_obs_1+1]
            col_2_data_1 = col_2_data[:n_obs_1+1]
            col_1_data_2 = col_1_data[n_obs_1+1:]
            col_2_data_2 = col_2_data[n_obs_1+1:]
            statistic_value = self.WaldTest_WithData(col_1_data_1, col_2_data_1, col_1_data_2, col_2_data_2)
            if statistic_value:
                tests_results.append(statistic_value)

        self.rellavent_data = self.word_counts_table
        if len(tests_results)==0:
            return ''     
        return self.update_p_value(sum(1 for num in tests_results if num <= t_value) / len(tests_results))


#### comparing diffrent columns dist ####

    def student_t_test(self, data_to_test, start_1, stop_1, start_2, stop_2):
        group_0 = data_to_test.iloc[start_1:stop_1+1, 0].tolist()
        group_1 = data_to_test.iloc[start_2:stop_2+1, 0].tolist()

        # Remove NaN and empty string values from group_0 and group_1
        group_0 = [x for x in group_0 if not (isinstance(x, str) and x == '') and not np.isnan(x)]
        group_1 = [x for x in group_1 if not (isinstance(x, str) and x == '') and not np.isnan(x)]

        student_t_test_result = ttest_ind(a=group_0, b=group_1, equal_var=True, alternative='less', nan_policy='omit')
        return student_t_test_result.pvalue
    
    def permutation_1_col(self, data_to_test, start_1, stop_1, start_2, stop_2):
        group_0 = data_to_test.iloc[start_1:stop_1+1, 0].tolist()
        group_1 = data_to_test.iloc[start_2:stop_2+1, 0].tolist()
        
        # Remove empty values from the lists
        group_0 = [val for val in group_0 if not np.isnan(val)]
        group_1 = [val for val in group_1 if not np.isnan(val)]

        # Check if both lists are empty and return np.nan
        if not group_0 or not group_1:
            return np.nan

        if np.sum(group_0 == np.nan) > 0 or np.sum(group_1 == np.nan) > 0:
            print("Something Fishy")
        p_value = permutation_test(data=(group_0, group_1), statistic=self.t_test_for_permutation, permutation_type='independent',
                                   vectorized=True, n_resamples=1000, alternative='less', axis=0)

        # Convert back to np.nan if p_value is zero-dimensional
        if np.isscalar(p_value):
            p_value = np.nan

        return p_value.pvalue

    def t_test_for_permutation(self, x_list, y_list, axis):
        x_array = np.array(x_list)
        y_array = np.array(y_list)

        if len(x_array) < 1 or len(y_array) < 1:
            return np.nan  # Return np.nan instead of ''

        _, p_value = ttest_ind(x_array, y_array, axis=axis, nan_policy='omit')
        return p_value

# update p_value (The extremity of the phenomenon is two-sided - on both sides of the normal distribution)
    def update_p_value(self, p_value):
        if p_value == '':
            return ''
        elif p_value> 0.5:
            return 1-p_value
        return p_value
        
##########################################################
# Unified pipeline for statistical tests between writers #
##########################################################  

# Getting the value of the t test and permutation for each test - as well as the number of successes for the entire pipeline
    def TestDifference(self, start_1, stop_1, start_2, stop_2, test_name ='', p_thresh = 0.1, var_test=True, row_test=True, count_corrections=True,  char_corrections='^'):
        print(test_name)
        n_success_t = 0
        n_success_p = 0
        self.rellavent_data = self.word_counts_table
        columns = self.rellavent_data.columns[:]

        # df to store results
        data = [np.nan] * len(self.table.columns) 
        new_row = pd.Series(data, index=self.table.columns, name=test_name + 't_test')
        self.table = pd.concat([self.table, pd.DataFrame([new_row])], ignore_index=True)
        new_row = pd.Series(data, index=self.table.columns, name=test_name + 'permutation_test')
        self.table = pd.concat([self.table, pd.DataFrame([new_row])], ignore_index=True)

        for i in range(0, len(self.word_counts_table.columns) , 2):
            col_1, col_2 = columns[i], columns[i+1]
            t_value = self.WaldTest(col_1, col_2, start_1, stop_1, start_2, stop_2)
            self.table.iloc[-2, i//2] = t_value
            if t_value or t_value == 0:
                p_value = self.PermuteWald(col_1, col_2, start_1, stop_1, start_2, stop_2, t_value, n_iter = 1000)
                self.table.iloc[-1, i//2] = p_value
                print("Column Names:", col_1, col_2)
                print("t_value: ", t_value)
                print("p_value: ", p_value)
                if t_value < p_thresh:
                    n_success_t += 1
                if p_value or p_value == 0:
                    if p_value < p_thresh:
                        n_success_p += 1

            self.table.rename(columns={self.table.columns[i//2]: col_1}, inplace=True)

# Additional tests if chekcing row's var or/and num of rows tests or/and num of corrections
        if var_test == True:       
            print("Row Variance Test:")
            self.compute_row_var()
            t_value = self.update_p_value(self.student_t_test(self.var_table, start_1, stop_1, start_2, stop_2))
            p_value = self.update_p_value(self.permutation_1_col(self.var_table, start_1, stop_1, start_2, stop_2))
            print("t_value: ", t_value)
            print("p_value: ", p_value)
            if t_value or t_value ==0:
                if t_value < p_thresh:
                    n_success_t += 1
            if p_value or p_value ==0:
                if p_value < p_thresh:
                    n_success_p += 1
            self.table.iloc[-2, i//2 + var_test] = t_value
            self.table.iloc[-1, i//2 + var_test] = p_value
            self.table.rename(columns={self.table.columns[i//2 +var_test]: 'row_variance'}, inplace=True)


        if row_test == True:
            print("Number of Rows Test:")
            self.compute_column_num_of_rows()
            t_value = self.update_p_value(self.student_t_test(self.num_of_rows_column, start_1, stop_1, start_2, stop_2))
            p_value = self.update_p_value(self.permutation_1_col(self.num_of_rows_column, start_1, stop_1, start_2, stop_2))
            print("t_value: ", t_value)
            print("p_value: ", p_value)
            if t_value or t_value ==0:
                if t_value < p_thresh:
                    n_success_t += 1
            if p_value or p_value ==0:
                if p_value < p_thresh:
                    n_success_p += 1
            self.table.iloc[-2, i//2 + var_test + row_test] = t_value
            self.table.iloc[-1, i//2 + var_test + row_test] = p_value
            self.table.rename(columns={self.table.columns[i//2 + var_test + row_test]: 'num_of_rows'}, inplace=True)

        
        if count_corrections == True:
            print("Number of Corrections Test:")
            self.count_corrections(char_corrections)
            t_value = self.update_p_value(self.student_t_test(self.corrections, start_1, stop_1, start_2, stop_2))
            p_value = self.update_p_value(self.permutation_1_col(self.corrections, start_1, stop_1, start_2, stop_2))
            if t_value or t_value ==0:
                if t_value < p_thresh:
                    n_success_t += 1
            if p_value or p_value ==0:
                if p_value < p_thresh:
                    n_success_p += 1
            self.table.iloc[-2, i//2 + var_test + row_test + count_corrections] = t_value
            self.table.iloc[-1, i//2 + var_test + row_test + count_corrections] = p_value
            self.table.rename(columns={self.table.columns[i//2 + var_test + row_test + count_corrections]: 'overall corrections'}, inplace=True)
            
            # self.export_table_statistical_test()
            return n_success_t, n_success_p

    def export_table_statistical_test(self):
        self.table.to_csv("Updated_Significance_Tests_Table.csv")      
#####################
# permutation tests #
#####################  

#  function
    def permutation_tests(self, start_1, stop_1, start_2, stop_2, type_of_test, bootstrap=0, test_name='', column_index_to_start=0, p_thresh = 0.1, num_iterations=10000, var_test=True, row_test=True, count_corrections=True):
        self.counting_successful_tests(start_1, stop_1, start_2, stop_2, type_of_test, bootstrap ,test_name , column_index_to_start, p_thresh, num_iterations, var_test, row_test, count_corrections)
        self.update_table_with_counts_proportion(test_name)
        self.export_table_permutaion_tests(test_name)

# define the rows to be tested and the transaction in wrtiers
    def counting_successful_tests(self, start_1, stop_1, start_2, stop_2, type_of_test, bootstrap, test_name, column_index_to_start, p_thresh, num_iterations, var_test, row_test, count_corrections):
        # shaffling the data 'num_iterations' times
        tests_results = []
        for j in range (num_iterations):
            print(j)
            counter = 0 
            indexes_to_shuffle = self.bootstrap_data(start_1, stop_1, start_2, stop_2) if bootstrap == 1 else self.preprocess_data(start_1, stop_2)
            # compute tests for each columns pair
            for i in range(column_index_to_start,len(self.rellavent_data.columns)-2,2):
                columns = self.rellavent_data.columns[column_index_to_start:]
                col_1, col_2 = columns[i], columns[i+1]
                if type_of_test == 't_test':
                     p_value = self.WaldTest(col_1, col_2, start_1, stop_1, start_2, stop_2)
                elif type_of_test == 'permutation':
                    t_value = self.WaldTest(col_1, col_2, start_1, stop_1, start_2, stop_2)
                    p_value = self.PermuteWald(col_1, col_2, start_1, stop_1, start_2, stop_2, t_value, n_iter = 1000)
                else: 
                    raise('An incorrect test type was entered in the function')
                p_value = self.update_p_value(p_value)
                if p_value == '':
                    continue 
                else:
                    if p_value <= p_thresh:
                        counter += 1

            # to var of rows and num of rows            
            if var_test == True:
                self.copy_var_table = self.var_table.copy()
                self.copy_var_table.iloc[np.arange(start_1-1, stop_2), :] = self.var_table.iloc[indexes_to_shuffle, :]
                if type_of_test == 't_test': 
                    p_value = self.update_p_value(self.student_t_test(self.copy_var_table, start_1, stop_1, start_2, stop_2))
                elif type_of_test == 'permutation':
                    p_value = self.update_p_value(self.permutation_1_col(self.copy_var_table, start_1, stop_1, start_2, stop_2))
                if p_value <= p_thresh:
                    counter += 1             
                
            if row_test == True:
                self.copy_num_of_rows_column = self.num_of_rows_column.copy()
                self.copy_num_of_rows_column.iloc[np.arange(start_1-1, stop_2), :] = self.num_of_rows_column.iloc[indexes_to_shuffle, :]
                if type_of_test == 't_test': 
                    p_value = self.update_p_value(self.student_t_test(self.copy_num_of_rows_column, start_1, stop_1, start_2, stop_2))
                elif type_of_test == 'permutation':
                    p_value = self.update_p_value(self.permutation_1_col(self.copy_num_of_rows_column, start_1, stop_1, start_2, stop_2))
                p_value = self.update_p_value(p_value)
                if p_value <= p_thresh:
                    counter += 1
                    
            if count_corrections == True:
                self.copy_corrections = self.corrections.copy()
                self.copy_corrections.iloc[np.arange(start_1-1, stop_2), :] = self.corrections.iloc[indexes_to_shuffle, :]
                if type_of_test == 't_test': 
                    p_value = self.update_p_value(self.student_t_test(self.copy_corrections, start_1, stop_1, start_2, stop_2))
                elif type_of_test == 'permutation':
                    p_value = self.update_p_value(self.permutation_1_col(self.copy_corrections, start_1, stop_1, start_2, stop_2))
                p_value = self.update_p_value(p_value)
                if p_value <= p_thresh:
                    counter += 1
                                    
            tests_results.append(counter)
        self.tests_results = tests_results
                  
# Filling up the tables
    def create_table_permutation_tests(self, length=20):
        row_names = range(length + 1)
        self.permutation_table = pd.DataFrame(index=range(length + 1))
            
    def update_table_with_counts_proportion(self, test_name):
        # add an empty row to self.table
        data = [np.nan] * len(self.permutation_table.index)
        new_column = pd.Series(data, index=self.permutation_table.index, name=test_name)
        self.permutation_table = pd.concat([self.permutation_table, new_column], axis=1)
        
        # Count the occurrences of each number in self.tests_results
        counts = {i: self.tests_results.count(i) for i in range(len(self.permutation_table.index))}

        list_length = len(self.tests_results)

        # Update the values in the specified row of the table
        for row_idx, count in counts.items():
            self.permutation_table.at[row_idx, test_name] = count / list_length

    def export_table_permutaion_tests(self, test_name):
        with pd.ExcelFile("C:/Users/yisha/.vscode/thesis/results/PERMUTATION_TESTS.xlsx") as writer:
             self.permutation_table.to_excel(writer, sheet_name=test_name, index=False)
                       
    def preprocess_data(self, start_row, stop_row, char_corrections='^'):
        # choozing rows for rellavant tests
        self.rellavent_data = self.word_counts_table.copy()
        indexes_to_shuffle = np.arange(start_row-1, stop_row)
        np.random.shuffle(indexes_to_shuffle)
        self.rellavent_data.iloc[np.arange(start_row-1, stop_row), :] = self.word_counts_table.iloc[indexes_to_shuffle, :]
        
        # Computing row var and row num
        self.compute_row_var()
        self.compute_column_num_of_rows()
        self.count_corrections(char_corrections)
        
        return indexes_to_shuffle

    def bootstrap_data(self, start_1, stop_1, start_2, stop_2, random_seed=None, num_samples=''):
        if not num_samples:
            num_samples = stop_2 - start_1 + 1
        if random_seed is not None:
            random.seed(random_seed)  # Set the random seed if provided

        # Sample row indexes within the specified range with replacement
        sampled_indexes = random.choices(range(start_1-1, stop_1), k=num_samples)

        # Create a new DataFrame from the sampled rows
        self.rellavent_data = self.word_counts_table.iloc[sampled_indexes]

        # Return the sampled row indexes (with duplicates)
        return sampled_indexes

#########################
#### Graths Plotting ####
#########################


    def plot_single_column_data(self, save_path=None, vertical_lines=None):
        david_font = FontProperties(family='David', size=30)
        arial_font = FontProperties(family='Arial', size=14)

        # Set font sizes
        plt.rcParams.update({
            'font.family': 'David',  # Font family
            'axes.titlesize': 32,  # Title size
            'axes.labelsize': 24,  # Axes labels size
            'xtick.labelsize': 26,  # X-axis tick size
            'ytick.labelsize': 26,  # Y-axis tick size
            'legend.fontsize': 26,  # Legend font size
            'legend.title_fontsize': 22  # Legend title font size
        })

        # Dictionary to map table to titles and y-axis labels
        table_info = {
            'var_table': ("Rows Variance per Column", "Rows Variance"),
            'num_of_rows_column': ("Number of Rows per Column", "Number of Rows"),
            'corrections': ("Number of Corrections per Column", "Number of Corrections")
        }

        for table_attr, (title, ylabel) in table_info.items():
            table = getattr(self, table_attr)

            # Prepare the data
            data = table.copy()
            data.columns = ['value']

            # Create a bar plot
            plt.figure(figsize=(12, 6))
            bars = plt.bar(data.index, data['value'], color='#86c5da', alpha=.9)

            # Apply dashed bars to 'num_of_rows_column' specifically
            if table_attr == 'num_of_rows_column':
                for i in [26, 53]:  # Specify indices for dashed bars
                    bars[i].set_hatch('//')
                    bars[i].set_edgecolor('black')

            # Set x-axis labels to show every 5th label
            plt.xticks(range(0, len(data.index), 5))

            # Set y-axis limits for specific plots
            if table_attr == 'num_of_rows_column':
                plt.ylim(26, data['value'].max() + 1)

            # Add vertical lines if provided
            if vertical_lines:
                for line in vertical_lines:
                    plt.axvline(x=line + 0.5, color='black', linestyle='-', linewidth=2)

            plt.title(title, fontproperties=david_font)
            plt.ylabel(ylabel, fontproperties=david_font)
            plt.xlabel(None)  # No x-label

            # Save the plot
            plt.tight_layout()
            if save_path:
                plt.savefig(f"{save_path}/{table_attr}_plot.png")
            else:
                plt.savefig(f"{table_attr}_plot.png")
            plt.close()
